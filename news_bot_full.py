#!/usr/bin/env python3
"""
雅虎财经新闻每日推送系统 (完整版)
通过RSS获取Yahoo Finance最新新闻，抓取全文并翻译为中文
确保至少包含：政策/宏观经济、行业动态、公司新闻各一条
"""

import json
import argparse
import requests
import feedparser
import time
import re
import os
import logging
from datetime import datetime, timedelta
from email.utils import parsedate_to_datetime
from bs4 import BeautifulSoup
from deep_translator import GoogleTranslator
from docx import Document
from docx.shared import Pt, RGBColor, Inches, Cm, Emu
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.section import WD_ORIENT
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml

# ─── 日志配置 ───
logging.basicConfig(level=logging.INFO, format='%(asctime)s - %(levelname)s - %(message)s')
logger = logging.getLogger(__name__)

# ─── 路径配置 ───
PROJECT_ROOT = os.path.dirname(os.path.abspath(__file__))
CONFIG_PATH = os.path.join(PROJECT_ROOT, 'push_config.json')
OUTPUT_DIR = os.path.join(PROJECT_ROOT, 'output')

# ─── 请求头 ───
HEADERS = {
    'User-Agent': 'Mozilla/5.0 (Windows NT 10.0; Win64; x64) AppleWebKit/537.36 '
                  '(KHTML, like Gecko) Chrome/120.0.0.0 Safari/537.36',
    'Accept-Language': 'en-US,en;q=0.9',
}

# ─── Yahoo Finance RSS 源 ───
RSS_FEEDS = [
    'https://finance.yahoo.com/news/rssindex',
    'https://finance.yahoo.com/rss/topstories',
]

# ─── 需要获取个股新闻的代码 ───
TICKERS = ['AAPL', 'MSFT', 'GOOGL', 'AMZN', 'TSLA', 'META', 'NVDA', 'JPM']
TICKER_RSS_TEMPLATE = 'https://finance.yahoo.com/rss/headline?s={ticker}'

# ─── 新闻分类关键词 ───
CATEGORY_KEYWORDS = {
    '政策/宏观经济': [
        'gdp', 'pce', 'inflation', 'fed', 'federal reserve', 'interest rate',
        'tariff', 'trade war', 'trade deal', 'policy', 'regulation', 'government',
        'biden', 'trump', 'congress', 'senate', 'legislation', 'white house',
        'treasury', 'fiscal', 'monetary', 'debt ceiling', 'deficit', 'surplus',
        'cpi', 'jobs report', 'unemployment', 'nonfarm', 'payroll',
        'economic data', 'recession', 'stimulus', 'sanctions', 'geopolitical',
        'china trade', 'eu regulation', 'central bank', 'rate cut', 'rate hike',
        'consumer spending', 'housing market', 'labor market',
    ],
    '行业动态': [
        'etf', 'sector', 'industry', 'automotive', 'ev', 'electric vehicle',
        'semiconductor', 'chip', 'ai', 'artificial intelligence', 'robotics',
        'energy', 'oil', 'gas', 'solar', 'wind', 'renewable',
        'healthcare', 'pharma', 'biotech', 'finance', 'banking', 'fintech',
        'real estate', 'retail', 'e-commerce', 'crypto', 'bitcoin', 'blockchain',
        'cloud computing', 'cybersecurity', 'streaming', '5g', 'autonomous',
        'space', 'quantum', 'data center', 'gigawatt',
    ],
    '公司新闻': [
        'earnings', 'revenue', 'profit', 'quarterly', 'annual report', 'guidance',
        'stock', 'shares', 'dividend', 'buyback', 'ipo', 'merger', 'acquisition',
        'ceo', 'executive', 'layoff', 'hiring', 'partnership', 'product launch',
        'apple', 'tesla', 'nvidia', 'amazon', 'google', 'alphabet', 'meta',
        'microsoft', 'berkshire', 'netflix', 'amd', 'intel', 'broadcom',
        'jpmorgan', 'goldman', 'walmart', 'costco', 'magnificent seven',
        'market cap', 'valuation', 'upgrade', 'downgrade', 'target price',
        'analyst', 'rating',
    ],
}

MIN_PER_CATEGORY = 1
MAX_ARTICLE_CHARS = 3000
TRANSLATE_CHUNK_SIZE = 4500
REQUEST_MAX_RETRIES = 3
REQUEST_RETRY_DELAY = 1.2

# ─── 广告 / 无关内容过滤规则 ───
# 包含任一关键词的行将被剔除（大小写不敏感）
AD_KEYWORDS_EN = [
    # 营销号召 / CTA
    'sign up', 'subscribe now', 'click here', 'download now', 'buy now',
    'limited time', 'free trial', 'get started', 'join now', 'act now',
    'don\'t miss', 'learn more', 'exclusive offer', 'special offer',
    'promo code', 'coupon', 'discount', 'watch now', 'explore the',
    'explore which', 'explore products',
    # 付费墙提示
    'continue reading', 'read more', 'read the full', 'full article',
    'unlock this', 'premium content', 'members only',
    # 免责 / 版权
    'all rights reserved', 'terms of use', 'privacy policy',
    'not financial advice', 'does not constitute', 'disclaimer',
    'copyright', '\u00a9',
    'is a trademark', 'are trademarks', 'trademark of',
    # 广告插入
    'sponsored', 'advertisement', 'promoted content', 'paid post',
    'partner content', 'presented by',
    # 投资推广
    'just launched', '/share', 'investors have joined',
    'round of funding', r'explore .* etf',
    # 网站导航 / 布局残留
    'related articles', 'trending now', 'you may also like',
    'recommended for you', 'most popular', 'see also',
    'in another article', 'check out our',
    'real-time prices', 'provided by lseg', 'provided by factset',
    'instant access to exclusive', 'get market trend',
    'education videos', 'webinars and stock',
    r'information .* is for informational and educational purposes',
]
AD_KEYWORDS_CN = [
    # 营销号召
    '立即注册', '立即订阅', '点击这里', '立即下载', '立即购买',
    '限时优惠', '免费试用', '立即开始', '立即加入', '不要错过',
    '了解更多', '独家优惠', '特别优惠', '促销代码',
    '探索哪种', '探索产品',
    # 付费墙
    '继续阅读', '阅读全文', '解锁此', '高级内容', '仅限会员',
    # 免责 / 版权
    '版权所有', '使用条款', '隐私政策', '不构成投资建议',
    '免责声明', '信息如有更改', '商标',
    # 广告
    '赞助内容', '广告', '推广内容', '合作内容',
    # 投资推广
    '刚刚启动了', '投资者已加入', '一轮融资',
    # 网站导航
    '相关文章', '热门推荐', '你可能还喜欢', '查看我们的',
    '另一篇文章', '在这里下载', '只需', '即可使用',
    '实时价格', '报价并非来自所有市场',
    '立即访问独家', '获取市场动态', '教育视频',
    '仅供参考和教育', '仅供信息和教育',
]


def _is_ad_line(line, lang='en'):
    """判断一行是否为广告/无关内容"""
    low = line.lower().strip()
    if not low:
        return True
    keywords = AD_KEYWORDS_EN if lang == 'en' else AD_KEYWORDS_CN
    for kw in keywords:
        if re.search(kw, low):
            return True
    return False


def _is_orphan_fragment(line, min_len=8):
    """判断是否为残留的空洞碎片（链接文字、按钮等）"""
    stripped = line.strip()
    if len(stripped) < min_len:
        return True
    # 纯特殊字符 / 数字 / 日期
    if re.match(r'^[\d\s/\-\.:,;|*#\u2022\u25cf]+$', stripped):
        return True
    return False


def clean_article_text(text, lang='en'):
    """清洗文章文本：剔除广告、导航、免责声明、碎片等"""
    if not text:
        return text

    lines = text.split('\n')
    cleaned = []

    for line in lines:
        stripped = line.strip()
        # 跳过广告行
        if _is_ad_line(stripped, lang):
            continue
        # 跳过无意义碎片
        if _is_orphan_fragment(stripped):
            continue
        cleaned.append(stripped)

    # 合并被错误断开的句子
    if lang == 'en':
        merged = []
        for line in cleaned:
            if (merged
                    and not merged[-1].endswith(('.', '!', '?', '"', '\u201d'))
                    and line and line[0].islower()):
                merged[-1] = merged[-1] + ' ' + line
            else:
                merged.append(line)
        cleaned = merged
    else:
        # 中文：上一行不以句号/问号/叹号/引号结尾 → 拼接
        merged = []
        for line in cleaned:
            if (merged
                    and not re.search(r'[。！？\u201d\u2019.!?]$', merged[-1])
                    and len(line) > 0):
                merged[-1] = merged[-1] + line
            else:
                merged.append(line)
        cleaned = merged

    return '\n'.join(cleaned)


def parse_published_time(published):
    """将RSS发布时间解析为datetime，失败时返回最小时间"""
    if not published:
        return datetime.min
    try:
        dt = parsedate_to_datetime(published)
        # 统一为不含时区的时间，便于排序比较
        if hasattr(dt, 'tzinfo') and dt.tzinfo is not None:
            return dt.replace(tzinfo=None)
        return dt
    except Exception:
        return datetime.min


def safe_get(url, headers=None, timeout=15, session=None, allow_redirects=True,
             max_retries=REQUEST_MAX_RETRIES):
    """带重试的GET请求，返回Response或None"""
    client = session if session is not None else requests
    for attempt in range(max_retries):
        try:
            resp = client.get(url, headers=headers, timeout=timeout, allow_redirects=allow_redirects)
            if resp.status_code == 200:
                return resp
            # 429/5xx视为可重试
            if resp.status_code == 429 or resp.status_code >= 500:
                logger.warning(f"请求失败(可重试) {resp.status_code}: {url}")
            else:
                logger.warning(f"请求失败(不重试) {resp.status_code}: {url}")
                return None
        except (requests.exceptions.Timeout, requests.exceptions.ConnectionError) as e:
            logger.warning(f"请求异常(尝试{attempt + 1}): {url} | {e}")
        except Exception as e:
            logger.warning(f"请求异常(终止): {url} | {e}")
            return None

        if attempt < max_retries - 1:
            time.sleep(REQUEST_RETRY_DELAY * (attempt + 1))

    return None


# ══════════════════════════════════════════════════════════════════════
# 翻译模块
# ══════════════════════════════════════════════════════════════════════

def _split_sentences(text):
    """将英文文本按句子边界切分，保留完整句子"""
    parts = re.split(r'(?<=[.!?])\s+', text)
    return [p.strip() for p in parts if p.strip()]


def translate_text(text, max_retries=3):
    """将英文文本翻译为中文，按句子边界分块，翻译后清洗格式"""
    if not text or not text.strip():
        return text

    # 如果已经主要是中文则跳过
    chinese_chars = len(re.findall(r'[\u4e00-\u9fff]', text))
    if chinese_chars > len(text) * 0.3:
        return text

    translator = GoogleTranslator(source='en', target='zh-CN')

    def _do_translate(t):
        for attempt in range(max_retries):
            try:
                result = translator.translate(t)
                return result if result else t
            except Exception as e:
                logger.warning(f"翻译失败(尝试{attempt + 1}): {e}")
                time.sleep(1)
        return t

    # 短文本直接翻译
    if len(text) <= TRANSLATE_CHUNK_SIZE:
        result = _do_translate(text)
        return clean_article_text(result, lang='cn')

    # 长文本：按句子边界分块，避免在句子中间截断
    sentences = _split_sentences(text)
    chunks = []
    current_chunk = ''

    for sent in sentences:
        if len(current_chunk) + len(sent) + 1 > TRANSLATE_CHUNK_SIZE:
            if current_chunk:
                chunks.append(current_chunk)
            current_chunk = sent
        else:
            current_chunk = current_chunk + ' ' + sent if current_chunk else sent
    if current_chunk:
        chunks.append(current_chunk)

    translated_parts = []
    for chunk in chunks:
        result = _do_translate(chunk)
        translated_parts.append(result)
        time.sleep(0.3)

    combined = '\n'.join(translated_parts)
    return clean_article_text(combined, lang='cn')


# ══════════════════════════════════════════════════════════════════════
# 新闻抓取模块
# ══════════════════════════════════════════════════════════════════════

def fetch_rss_entries():
    """从多个Yahoo Finance RSS源获取新闻条目"""
    all_entries = []
    seen_links = set()
    seen_titles = set()

    def _normalize_title(title):
        return re.sub(r'\s+', ' ', (title or '').strip().lower())

    # 1) 综合RSS源
    for rss_url in RSS_FEEDS:
        try:
            logger.info(f"正在获取RSS: {rss_url}")
            resp = safe_get(rss_url, headers=HEADERS, timeout=15)
            if resp is not None:
                feed = feedparser.parse(resp.text)
                for entry in feed.entries:
                    link = entry.get('link', '')
                    title = entry.get('title', '')
                    title_key = _normalize_title(title)
                    if link and link not in seen_links and title_key not in seen_titles:
                        seen_links.add(link)
                        seen_titles.add(title_key)
                        # 提取RSS自带的摘要/描述
                        rss_summary = entry.get('summary', '') or entry.get('description', '')
                        # 清洗 HTML 标签
                        if rss_summary:
                            rss_summary = BeautifulSoup(rss_summary, 'html.parser').get_text(strip=True)
                        all_entries.append({
                            'title': title,
                            'link': link,
                            'published': entry.get('published', ''),
                            'published_dt': parse_published_time(entry.get('published', '')),
                            'source': 'Yahoo Finance',
                            'ticker': '',
                            'rss_summary': rss_summary,
                        })
        except Exception as e:
            logger.error(f"RSS获取异常 {rss_url}: {e}")

    # 2) 个股RSS源
    for ticker in TICKERS:
        rss_url = TICKER_RSS_TEMPLATE.format(ticker=ticker)
        try:
            resp = safe_get(rss_url, headers=HEADERS, timeout=10)
            if resp is not None:
                feed = feedparser.parse(resp.text)
                for entry in feed.entries[:5]:
                    link = entry.get('link', '')
                    title = entry.get('title', '')
                    title_key = _normalize_title(title)
                    if link and link not in seen_links and title_key not in seen_titles:
                        seen_links.add(link)
                        seen_titles.add(title_key)
                        rss_summary = entry.get('summary', '') or entry.get('description', '')
                        if rss_summary:
                            rss_summary = BeautifulSoup(rss_summary, 'html.parser').get_text(strip=True)
                        all_entries.append({
                            'title': title,
                            'link': link,
                            'published': entry.get('published', ''),
                            'published_dt': parse_published_time(entry.get('published', '')),
                            'source': 'Yahoo Finance',
                            'ticker': ticker,
                            'rss_summary': rss_summary,
                        })
        except Exception as e:
            logger.warning(f"个股RSS获取失败 {ticker}: {e}")

    # 按发布时间从新到旧排序
    all_entries.sort(key=lambda x: x.get('published_dt', datetime.min), reverse=True)

    logger.info(f"RSS共获取 {len(all_entries)} 条不重复新闻")
    return all_entries


def scrape_article_body(url, rss_summary=''):
    """抓取文章正文内容，多策略尝试，rss_summary作为最终备选"""
    session = requests.Session()
    session.headers.update(HEADERS)
    # 设置常见cookie以绕过Yahoo consent页面
    session.cookies.set('A1', 'v=1', domain='.yahoo.com')
    session.cookies.set('GUC', 'AQABCAFn', domain='.yahoo.com')

    text = ''
    try:
        resp = safe_get(url, timeout=20, session=session, allow_redirects=True)
        if resp is None:
            return rss_summary[:MAX_ARTICLE_CHARS] if rss_summary else ''

        soup = BeautifulSoup(resp.text, 'html.parser')

        # ── 策略1: Yahoo Finance caas-body (最常见的正文容器) ──
        body = soup.find('div', {'class': lambda c: c and 'caas-body' in c})

        # ── 策略2: 按data属性查找正文区域 ──
        if not body:
            body = soup.find('div', attrs={'data-test-locator': 'articleBody'})
        if not body:
            body = soup.find('div', class_='body')

        # ── 策略3: article 标签 ──
        if not body:
            body = soup.find('article')

        # ── 策略4: 查找包含大量 <p> 的最大 div ──
        if not body:
            candidates = soup.find_all('div')
            best, best_len = None, 0
            for div in candidates:
                ps = div.find_all('p', recursive=True)
                total = sum(len(p.get_text(strip=True)) for p in ps)
                if total > best_len:
                    best_len = total
                    best = div
            if best and best_len > 200:
                body = best

        if body:
            # 去除无关标签
            for tag in body.find_all(['script', 'style', 'nav', 'footer', 'aside',
                                       'figure', 'iframe', 'button', 'svg',
                                       'form', 'input', 'select', 'label']):
                tag.decompose()
            # 移除广告相关的 div（class/id 含 ad、promo、sidebar、related 等）
            for tag in body.find_all(['div', 'section', 'aside']):
                cls = ' '.join(tag.get('class', []) or []).lower()
                tid = (tag.get('id') or '').lower()
                combined = cls + ' ' + tid
                if re.search(r'ad[-_]?|promo|sponsor|sidebar|related|newsletter|signup|banner|widget|outbrain|taboola', combined):
                    tag.decompose()
            # 移除链接文字残留（独立的 <a> 标签且不在 <p> 内）
            for a_tag in body.find_all('a'):
                parent_p = a_tag.find_parent('p')
                if not parent_p and len(a_tag.get_text(strip=True)) < 50:
                    a_tag.decompose()
            text = body.get_text('\n', strip=True)

        # ── 策略5: 退化 — 收集所有有意义的 <p> ──
        if len(text) < 100:
            paragraphs = soup.find_all('p')
            p_text = '\n'.join(
                p.get_text(strip=True) for p in paragraphs
                if len(p.get_text(strip=True)) > 30
            )
            if len(p_text) > len(text):
                text = p_text

        # ── 策略6: 从 meta 标签获取描述 ──
        if len(text) < 80:
            meta_desc = (
                soup.find('meta', attrs={'name': 'description'})
                or soup.find('meta', attrs={'property': 'og:description'})
            )
            if meta_desc:
                desc = meta_desc.get('content', '')
                if desc and len(desc) > len(text):
                    text = desc

        # ── 策略7: 使用 RSS 自带的摘要作为最终备选 ──
        if len(text) < 50 and rss_summary:
            text = rss_summary

        # 清洗英文原文：剔除广告、碎片、合并断句
        text = clean_article_text(text, lang='en')

        return text[:MAX_ARTICLE_CHARS] if text else ''

    except requests.exceptions.Timeout:
        logger.warning(f"抓取超时 {url}")
    except requests.exceptions.ConnectionError:
        logger.warning(f"连接失败 {url}")
    except Exception as e:
        logger.warning(f"抓取文章失败 {url}: {e}")

    # 网络层失败时仍可使用RSS摘要
    return rss_summary[:MAX_ARTICLE_CHARS] if rss_summary else ''


# ══════════════════════════════════════════════════════════════════════
# 分类模块
# ══════════════════════════════════════════════════════════════════════

def classify_news(title, body=''):
    """根据标题和正文分类新闻"""
    text = (title + ' ' + body).lower()

    scores = {}
    for category, keywords in CATEGORY_KEYWORDS.items():
        score = sum(1 for kw in keywords if kw in text)
        scores[category] = score

    if max(scores.values()) == 0:
        return '其他'
    return max(scores, key=scores.get)


def ensure_min_per_category(news_list):
    """确保每个必需分类至少有 MIN_PER_CATEGORY 条新闻"""
    category_counts = {}
    for item in news_list:
        cat = item.get('category', '其他')
        category_counts[cat] = category_counts.get(cat, 0) + 1

    required_cats = ['政策/宏观经济', '行业动态', '公司新闻']
    missing = [c for c in required_cats if category_counts.get(c, 0) < MIN_PER_CATEGORY]

    if not missing:
        return news_list

    logger.info(f"以下分类新闻不足，尝试补充: {missing}")

    # 尝试将 "其他" 类新闻重新分配
    for item in news_list:
        if item['category'] == '其他' and missing:
            text = (item.get('title_en', '') + ' ' + item.get('body_en', '')).lower()
            for cat in missing[:]:
                kws = CATEGORY_KEYWORDS[cat]
                if any(kw in text for kw in kws):
                    item['category'] = cat
                    category_counts[cat] = category_counts.get(cat, 0) + 1
                    if category_counts[cat] >= MIN_PER_CATEGORY:
                        missing.remove(cat)
                    break

    # 如果还有缺失，强制将最后几条 "其他" 分配
    for item in news_list:
        if not missing:
            break
        if item['category'] == '其他':
            item['category'] = missing.pop(0)

    return news_list


# ══════════════════════════════════════════════════════════════════════
# 推送模块
# ══════════════════════════════════════════════════════════════════════

class NewsPusher:
    """新闻推送器"""

    def __init__(self, config_file=None):
        self.config = {}
        config_paths = [config_file, CONFIG_PATH]
        for path in config_paths:
            if path and os.path.exists(path):
                try:
                    with open(path, 'r', encoding='utf-8') as f:
                        self.config = json.load(f)
                    logger.info(f"推送配置已加载: {path}")
                    break
                except Exception as e:
                    logger.warning(f"配置加载失败 {path}: {e}")

    def push_wechat_work(self, message, webhook_url=None):
        url = webhook_url or self.config.get('wechat_work_webhook')
        if not url or 'YOUR' in url:
            return False
        try:
            r = requests.post(url, json={"msgtype": "text", "text": {"content": message}}, timeout=30)
            return r.json().get('errcode') == 0
        except Exception:
            return False

    def push_dingtalk(self, message, webhook_url=None):
        url = webhook_url or self.config.get('dingtalk_webhook')
        if not url or 'YOUR' in url:
            return False
        try:
            r = requests.post(url, json={"msgtype": "text", "text": {"content": message}}, timeout=30)
            return r.json().get('errcode') == 0
        except Exception:
            return False

    def push_slack(self, message, webhook_url=None):
        url = webhook_url or self.config.get('slack_webhook')
        if not url or 'YOUR' in url:
            return False
        try:
            r = requests.post(url, json={"text": message}, timeout=30)
            return r.status_code == 200
        except Exception:
            return False

    def push_telegram(self, message, bot_token=None, chat_id=None):
        token = bot_token or self.config.get('telegram_bot_token')
        chat = chat_id or self.config.get('telegram_chat_id')
        if not token or not chat or 'YOUR' in token:
            return False
        try:
            url = f"https://api.telegram.org/bot{token}/sendMessage"
            r = requests.post(url, json={"chat_id": chat, "text": message, "parse_mode": "HTML"}, timeout=30)
            return r.json().get('ok')
        except Exception:
            return False

    def push_all(self, message):
        results = {}
        if self.config.get('wechat_work_webhook'):
            results['wechat_work'] = self.push_wechat_work(message)
        if self.config.get('dingtalk_webhook'):
            results['dingtalk'] = self.push_dingtalk(message)
        if self.config.get('slack_webhook'):
            results['slack'] = self.push_slack(message)
        if self.config.get('telegram_bot_token'):
            results['telegram'] = self.push_telegram(message)
        return results


# ══════════════════════════════════════════════════════════════════════
# 报告生成模块
# ══════════════════════════════════════════════════════════════════════

def format_report(news_list):
    """生成中文全文报告"""
    date_str = datetime.now().strftime('%Y-%m-%d')
    total = len(news_list)

    lines = []
    lines.append('=' * 70)
    lines.append('📰  雅虎财经新闻中文全文报告')
    lines.append(f'报告日期: {date_str}')
    lines.append(f'新闻总数: {total} 条')
    lines.append(f'生成时间: {datetime.now().strftime("%Y-%m-%d %H:%M:%S")}')
    lines.append('=' * 70)

    icon_map = {
        '政策/宏观经济': '📊', '行业动态': '🏭',
        '公司新闻': '🏢', '其他': '📌',
    }

    for category in ['政策/宏观经济', '行业动态', '公司新闻', '其他']:
        cat_news = [n for n in news_list if n.get('category') == category]
        if not cat_news:
            continue

        icon = icon_map.get(category, '📝')
        lines.append('')
        lines.append('─' * 70)
        lines.append(f'{icon}  【{category}】 — 共 {len(cat_news)} 条')
        lines.append('─' * 70)

        for i, item in enumerate(cat_news, 1):
            lines.append('')
            lines.append(f'  [{i}] {item.get("title_cn", item.get("title", ""))}')
            if item.get('ticker'):
                lines.append(f'      相关股票: {item["ticker"]}')
            lines.append(f'      来源: {item.get("source", "Yahoo Finance")}')
            lines.append(f'      发布: {item.get("published", "N/A")}')
            lines.append(f'      链接: {item.get("link", "")}')
            lines.append('')
            body_cn = item.get('body_cn', '')
            if body_cn:
                lines.append('      ── 全文翻译 ──')
                for para in body_cn.split('\n'):
                    para = para.strip()
                    if para and len(para) > 5:
                        # 跳过翻译后残留的广告行
                        if _is_ad_line(para, lang='cn'):
                            continue
                        lines.append(f'      {para}')
            else:
                lines.append('      (正文获取失败)')
            lines.append('')

    lines.append('=' * 70)
    lines.append('数据来源: Yahoo Finance RSS + 网页抓取')
    lines.append('翻译引擎: Google Translate')
    lines.append('=' * 70)

    return '\n'.join(lines)


def _make_brief(item, max_len=120):
    """从正文或RSS摘要中提取简短描述"""
    # 优先使用翻译后的正文
    text = item.get('body_cn', '') or item.get('body_en', '') or item.get('rss_summary', '')
    if not text:
        return ''
    # 取第一段有意义的内容
    for line in text.split('\n'):
        line = line.strip()
        if len(line) > 15:
            if len(line) > max_len:
                return line[:max_len] + '...'
            return line
    # 如果没有足够长的段落，直接截断
    text = text.replace('\n', ' ').strip()
    if len(text) > max_len:
        return text[:max_len] + '...'
    return text


def format_summary(news_list):
    """生成带摘要的推送消息（标题+内容概要）"""
    date_str = datetime.now().strftime('%Y-%m-%d')
    total = len(news_list)
    icon_map = {
        '政策/宏观经济': '📊', '行业动态': '🏭',
        '公司新闻': '🏢', '其他': '📌',
    }

    lines = []
    lines.append(f'📰 雅虎财经新闻摘要 ({date_str}) — 共{total}条')
    lines.append('=' * 50)

    for category in ['政策/宏观经济', '行业动态', '公司新闻', '其他']:
        cat_news = [n for n in news_list if n.get('category') == category]
        if not cat_news:
            continue
        icon = icon_map.get(category, '📝')
        lines.append(f'\n{icon} 【{category}】({len(cat_news)}条)')
        for i, item in enumerate(cat_news[:5], 1):
            title = item.get('title_cn', item.get('title', ''))
            lines.append(f'  {i}. {title}')
            brief = _make_brief(item)
            if brief:
                lines.append(f'     ▸ {brief}')

    lines.append(f'\n数据来源: Yahoo Finance | {datetime.now().strftime("%H:%M")}')
    return '\n'.join(lines)


def _set_cell_font(cell, font_name='微软雅黑', font_size=Pt(10), bold=False, color=None):
    """设置单元格内所有段落的字体"""
    for paragraph in cell.paragraphs:
        for run in paragraph.runs:
            run.font.name = font_name
            run._element.rPr.rFonts.set(qn('w:eastAsia'), font_name)
            run.font.size = font_size
            run.font.bold = bold
            if color:
                run.font.color.rgb = color


def _set_run_font(run, font_name='微软雅黑', font_size=Pt(10.5), bold=False, color=None):
    """设置 run 的中英文字体"""
    run.font.name = font_name
    run._element.rPr.rFonts.set(qn('w:eastAsia'), font_name)
    run.font.size = font_size
    run.font.bold = bold
    if color:
        run.font.color.rgb = color


def _add_horizontal_line(doc, color='4472C4'):
    """在文档中添加一条水平分隔线"""
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(4)
    pPr = p._element.get_or_add_pPr()
    pBdr = parse_xml(
        f'<w:pBdr {nsdecls("w")}>'
        f'  <w:bottom w:val="single" w:sz="6" w:space="1" w:color="{color}"/>'
        f'</w:pBdr>'
    )
    pPr.append(pBdr)


def _shade_cells(cells, color='D9E2F3'):
    """为单元格设置背景色"""
    for cell in cells:
        shading = parse_xml(
            f'<w:shd {nsdecls("w")} w:fill="{color}" w:val="clear"/>'
        )
        cell._element.get_or_add_tcPr().append(shading)


def save_report_to_word(news_list, output_dir=None):
    """将新闻报告保存为排版精美的Word文档"""
    if output_dir is None:
        output_dir = OUTPUT_DIR

    os.makedirs(output_dir, exist_ok=True)
    history_dir = os.path.join(output_dir, 'history')
    os.makedirs(history_dir, exist_ok=True)

    date_tag = datetime.now().strftime('%Y%m%d')
    date_str = datetime.now().strftime('%Y-%m-%d')

    # ── 创建文档 & 页面设置 ──
    doc = Document()

    # 页面边距
    for section in doc.sections:
        section.top_margin = Cm(2.5)
        section.bottom_margin = Cm(2.5)
        section.left_margin = Cm(2.8)
        section.right_margin = Cm(2.8)

    # ── 设置全局默认字体（Normal 样式）──
    style_normal = doc.styles['Normal']
    style_normal.font.name = '微软雅黑'
    style_normal._element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')
    style_normal.font.size = Pt(10.5)
    style_normal.font.color.rgb = RGBColor(0x33, 0x33, 0x33)
    style_normal.paragraph_format.line_spacing = 1.5
    style_normal.paragraph_format.space_after = Pt(6)

    # ── 自定义 Heading 样式字体 ──
    for level in range(4):
        style_name = f'Heading {level + 1}' if level > 0 else 'Title'
        try:
            h_style = doc.styles[style_name] if level > 0 else doc.styles['Title']
        except KeyError:
            h_style = doc.styles[f'Heading {level + 1}']
        h_style.font.name = '微软雅黑'
        h_style._element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')

    # ═══════════════════════════════════════════
    # 封面区域
    # ═══════════════════════════════════════════
    doc.add_paragraph()  # 顶部留白

    # 主标题
    title_para = doc.add_paragraph()
    title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title_para.paragraph_format.space_after = Pt(4)
    run_title = title_para.add_run('雅虎财经新闻日报')
    _set_run_font(run_title, font_name='微软雅黑', font_size=Pt(26), bold=True,
                  color=RGBColor(0x1F, 0x3A, 0x5F))

    # 副标题
    sub_para = doc.add_paragraph()
    sub_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sub_para.paragraph_format.space_after = Pt(20)
    run_sub = sub_para.add_run('Yahoo Finance Daily News Report')
    _set_run_font(run_sub, font_name='Calibri', font_size=Pt(13), bold=False,
                  color=RGBColor(0x66, 0x66, 0x66))

    _add_horizontal_line(doc, color='1F3A5F')

    # 报告元信息（用表格排列，无边框）
    info_table = doc.add_table(rows=3, cols=2)
    info_table.alignment = WD_TABLE_ALIGNMENT.CENTER
    info_table.autofit = True

    info_data = [
        ('报告日期', date_str),
        ('生成时间', datetime.now().strftime('%Y-%m-%d %H:%M:%S')),
        ('新闻总数', f'{len(news_list)} 条'),
    ]
    for idx, (label, value) in enumerate(info_data):
        label_cell = info_table.rows[idx].cells[0]
        value_cell = info_table.rows[idx].cells[1]
        label_cell.text = label
        value_cell.text = value
        _set_cell_font(label_cell, font_size=Pt(11), bold=True, color=RGBColor(0x1F, 0x3A, 0x5F))
        _set_cell_font(value_cell, font_size=Pt(11))
        # 设置列宽
        label_cell.width = Cm(3)
        value_cell.width = Cm(8)

    # 去除信息表格边框
    for row in info_table.rows:
        for cell in row.cells:
            tc = cell._element
            tcPr = tc.get_or_add_tcPr()
            tcBorders = parse_xml(
                f'<w:tcBorders {nsdecls("w")}>'
                '  <w:top w:val="none" w:sz="0" w:space="0" w:color="auto"/>'
                '  <w:left w:val="none" w:sz="0" w:space="0" w:color="auto"/>'
                '  <w:bottom w:val="none" w:sz="0" w:space="0" w:color="auto"/>'
                '  <w:right w:val="none" w:sz="0" w:space="0" w:color="auto"/>'
                '</w:tcBorders>'
            )
            tcPr.append(tcBorders)

    doc.add_paragraph()  # 留白

    # ── 分类统计概览 ──
    cat_data = []
    for category in ['政策/宏观经济', '行业动态', '公司新闻', '其他']:
        cnt = len([n for n in news_list if n.get('category') == category])
        if cnt > 0:
            cat_data.append((category, cnt))

    if cat_data:
        overview_para = doc.add_paragraph()
        overview_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        overview_para.paragraph_format.space_after = Pt(8)
        run_ov = overview_para.add_run('分类概览')
        _set_run_font(run_ov, font_size=Pt(13), bold=True, color=RGBColor(0x1F, 0x3A, 0x5F))

        ov_table = doc.add_table(rows=1, cols=len(cat_data))
        ov_table.alignment = WD_TABLE_ALIGNMENT.CENTER
        cat_colors = {
            '政策/宏观经济': 'D6E4F0',
            '行业动态': 'E2EFDA',
            '公司新闻': 'FFF2CC',
            '其他': 'EDEDED',
        }
        for idx, (cat, cnt) in enumerate(cat_data):
            cell = ov_table.rows[0].cells[idx]
            cell.text = f'{cat}\n{cnt} 条'
            for p in cell.paragraphs:
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            _set_cell_font(cell, font_size=Pt(10), bold=True,
                           color=RGBColor(0x33, 0x33, 0x33))
            _shade_cells([cell], color=cat_colors.get(cat, 'EDEDED'))

    # ═══════════════════════════════════════════
    # 分页：正文内容
    # ═══════════════════════════════════════════
    doc.add_page_break()

    category_colors = {
        '政策/宏观经济': RGBColor(0x1F, 0x4E, 0x79),
        '行业动态': RGBColor(0x2E, 0x75, 0x2E),
        '公司新闻': RGBColor(0xBF, 0x8F, 0x00),
        '其他': RGBColor(0x59, 0x59, 0x59),
    }
    category_bar_colors = {
        '政策/宏观经济': '1F4E79',
        '行业动态': '2E752E',
        '公司新闻': 'BF8F00',
        '其他': '595959',
    }

    for category in ['政策/宏观经济', '行业动态', '公司新闻', '其他']:
        cat_news = [n for n in news_list if n.get('category') == category]
        if not cat_news:
            continue

        cat_color = category_colors.get(category, RGBColor(0x33, 0x33, 0x33))
        bar_color = category_bar_colors.get(category, '333333')

        # ── 分类标题 ──
        _add_horizontal_line(doc, color=bar_color)
        cat_heading = doc.add_paragraph()
        cat_heading.paragraph_format.space_before = Pt(6)
        cat_heading.paragraph_format.space_after = Pt(10)
        run_cat = cat_heading.add_run(f'【{category}】— 共 {len(cat_news)} 条')
        _set_run_font(run_cat, font_size=Pt(16), bold=True, color=cat_color)

        for i, item in enumerate(cat_news, 1):
            # ── 新闻标题 ──
            news_title_para = doc.add_paragraph()
            news_title_para.paragraph_format.space_before = Pt(10)
            news_title_para.paragraph_format.space_after = Pt(6)
            run_num = news_title_para.add_run(f'{i}. ')
            _set_run_font(run_num, font_size=Pt(13), bold=True, color=cat_color)
            run_news_title = news_title_para.add_run(
                item.get('title_cn', item.get('title', ''))
            )
            _set_run_font(run_news_title, font_size=Pt(13), bold=True,
                          color=RGBColor(0x1F, 0x1F, 0x1F))

            # ── 元信息：紧凑的灰色小字 ──
            meta_parts = []
            if item.get('ticker'):
                meta_parts.append(f'股票: {item["ticker"]}')
            meta_parts.append(f'来源: {item.get("source", "Yahoo Finance")}')
            if item.get('published'):
                meta_parts.append(f'发布: {item["published"]}')

            meta_para = doc.add_paragraph()
            meta_para.paragraph_format.space_before = Pt(0)
            meta_para.paragraph_format.space_after = Pt(2)
            run_meta = meta_para.add_run('  |  '.join(meta_parts))
            _set_run_font(run_meta, font_size=Pt(9), color=RGBColor(0x88, 0x88, 0x88))

            # 原文链接（单独一行，可点击样式）
            if item.get('link'):
                link_para = doc.add_paragraph()
                link_para.paragraph_format.space_before = Pt(0)
                link_para.paragraph_format.space_after = Pt(6)
                run_link_label = link_para.add_run('原文: ')
                _set_run_font(run_link_label, font_size=Pt(9),
                              color=RGBColor(0x88, 0x88, 0x88))
                run_link = link_para.add_run(item['link'])
                _set_run_font(run_link, font_name='Calibri', font_size=Pt(8.5),
                              color=RGBColor(0x05, 0x63, 0xC1))

            # ── 全文翻译 ──
            body_cn = item.get('body_cn', '')
            if body_cn:
                # "全文翻译" 小标签
                label_para = doc.add_paragraph()
                label_para.paragraph_format.space_before = Pt(6)
                label_para.paragraph_format.space_after = Pt(4)
                run_label = label_para.add_run('▎全文翻译')
                _set_run_font(run_label, font_size=Pt(10), bold=True, color=cat_color)

                for para_text in body_cn.split('\n'):
                    para_text = para_text.strip()
                    if para_text and len(para_text) > 5:
                        if _is_ad_line(para_text, lang='cn'):
                            continue
                        p = doc.add_paragraph()
                        p.paragraph_format.first_line_indent = Cm(0.8)
                        p.paragraph_format.space_after = Pt(4)
                        p.paragraph_format.line_spacing = 1.6
                        run_body = p.add_run(para_text)
                        _set_run_font(run_body, font_size=Pt(10.5),
                                      color=RGBColor(0x33, 0x33, 0x33))
            else:
                no_body = doc.add_paragraph()
                no_body.paragraph_format.space_after = Pt(6)
                run_no = no_body.add_run('（正文获取失败）')
                _set_run_font(run_no, font_size=Pt(10),
                              color=RGBColor(0xAA, 0xAA, 0xAA))

            # 新闻之间的淡色分隔线
            if i < len(cat_news):
                sep_para = doc.add_paragraph()
                sep_para.paragraph_format.space_before = Pt(6)
                sep_para.paragraph_format.space_after = Pt(6)
                pPr = sep_para._element.get_or_add_pPr()
                pBdr = parse_xml(
                    f'<w:pBdr {nsdecls("w")}>'
                    '  <w:bottom w:val="dotted" w:sz="4" w:space="1" w:color="CCCCCC"/>'
                    '</w:pBdr>'
                )
                pPr.append(pBdr)

    # ═══════════════════════════════════════════
    # 页脚 / 声明
    # ═══════════════════════════════════════════
    _add_horizontal_line(doc, color='1F3A5F')

    footer_para = doc.add_paragraph()
    footer_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    footer_para.paragraph_format.space_before = Pt(12)
    footer_para.paragraph_format.space_after = Pt(4)
    run_f1 = footer_para.add_run('数据来源: Yahoo Finance RSS + 网页抓取')
    _set_run_font(run_f1, font_size=Pt(9), color=RGBColor(0x88, 0x88, 0x88))

    footer_para2 = doc.add_paragraph()
    footer_para2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run_f2 = footer_para2.add_run('翻译引擎: Google Translate  |  自动生成，仅供参考')
    _set_run_font(run_f2, font_size=Pt(9), color=RGBColor(0x88, 0x88, 0x88))

    # ── 保存文件 ──
    doc_path = os.path.join(history_dir, f'news_report_{date_tag}.docx')
    doc.save(doc_path)
    logger.info(f"Word文档已保存: {doc_path}")

    latest_doc_path = os.path.join(output_dir, 'latest_report.docx')
    doc.save(latest_doc_path)
    logger.info(f"最新Word文档已保存: {latest_doc_path}")

    print(f"  📄 Word文档: {doc_path}")
    print(f"  📄 最新报告: {latest_doc_path}")

    return doc_path


def save_report(news_list, output_dir=None, save_text_json=False):
    """保存报告到文件"""
    if output_dir is None:
        output_dir = OUTPUT_DIR

    os.makedirs(output_dir, exist_ok=True)
    history_dir = os.path.join(output_dir, 'history')
    os.makedirs(history_dir, exist_ok=True)

    date_tag = datetime.now().strftime('%Y%m%d')

    if save_text_json:
        # 构建 JSON 数据（去除不需要序列化的大字段）
        serializable = []
        for item in news_list:
            serializable.append({
                'title_cn': item.get('title_cn', ''),
                'title_en': item.get('title_en', ''),
                'category': item.get('category', ''),
                'ticker': item.get('ticker', ''),
                'source': item.get('source', ''),
                'published': item.get('published', ''),
                'link': item.get('link', ''),
                'body_cn': item.get('body_cn', ''),
            })

        json_data = {
            'date': datetime.now().strftime('%Y-%m-%d'),
            'generated_at': datetime.now().strftime('%Y-%m-%d %H:%M:%S'),
            'total_news': len(news_list),
            'categories': {},
            'all_news': serializable,
        }
        for category in ['政策/宏观经济', '行业动态', '公司新闻', '其他']:
            cat_items = [n for n in news_list if n.get('category') == category]
            if cat_items:
                json_data['categories'][category] = {
                    'count': len(cat_items),
                    'headlines': [n.get('title_cn', '') for n in cat_items],
                }

        # 保存到 history
        json_path = os.path.join(history_dir, f'news_summary_{date_tag}.json')
        with open(json_path, 'w', encoding='utf-8') as f:
            json.dump(json_data, f, ensure_ascii=False, indent=2)
        logger.info(f"JSON 已保存: {json_path}")

        detail_path = os.path.join(history_dir, f'news_detail_{date_tag}.txt')
        with open(detail_path, 'w', encoding='utf-8') as f:
            f.write(format_report(news_list))
        logger.info(f"详细报告已保存: {detail_path}")

        summary_path = os.path.join(history_dir, f'news_summary_{date_tag}.txt')
        with open(summary_path, 'w', encoding='utf-8') as f:
            f.write(format_summary(news_list))
        logger.info(f"摘要已保存: {summary_path}")

        # 保存 latest
        with open(os.path.join(output_dir, 'latest_summary.json'), 'w', encoding='utf-8') as f:
            json.dump(json_data, f, ensure_ascii=False, indent=2)
        with open(os.path.join(output_dir, 'latest_detail.txt'), 'w', encoding='utf-8') as f:
            f.write(format_report(news_list))
        with open(os.path.join(output_dir, 'latest_summary.txt'), 'w', encoding='utf-8') as f:
            f.write(format_summary(news_list))
        logger.info(f"最新文本/JSON报告已保存到 {output_dir}")
    else:
        logger.info("已启用精简输出模式: 仅保存Word文档")

    # 保存Word文档
    save_report_to_word(news_list, output_dir)


# ══════════════════════════════════════════════════════════════════════
# 主流程
# ══════════════════════════════════════════════════════════════════════

def main():
    parser = argparse.ArgumentParser(description='雅虎财经新闻推送系统')
    parser.add_argument('--max-per-category', type=int, default=5, help='每个分类最多抓取条数')
    parser.add_argument('--max-total', type=int, default=20, help='总抓取条数上限')
    parser.add_argument('--output-dir', type=str, default=OUTPUT_DIR, help='输出目录')
    parser.add_argument('--report-format', choices=['word', 'all'], default='word',
                        help='报告输出格式: word(仅Word) 或 all(Word+TXT+JSON)')
    parser.add_argument('--no-push', action='store_true', help='仅生成报告，不推送')
    args = parser.parse_args()

    if args.max_per_category <= 0:
        args.max_per_category = 5
    if args.max_total <= 0:
        args.max_total = 20

    print(f"\n{'=' * 60}")
    print(f"雅虎财经新闻推送系统 — {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}")
    print(f"{'=' * 60}\n")

    # ── 第1步：获取RSS新闻列表 ──
    print("▶ 第1步: 从Yahoo Finance RSS获取新闻列表...")
    rss_entries = fetch_rss_entries()
    if not rss_entries:
        print("❌ 未获取到任何RSS新闻条目，程序退出")
        return
    print(f"  获取到 {len(rss_entries)} 条新闻条目")

    # ── 第2步：分类 & 选择新闻 ──
    print("\n▶ 第2步: 对新闻进行分类...")
    for entry in rss_entries:
        entry['category'] = classify_news(entry['title'], entry.get('rss_summary', ''))
        entry['title_en'] = entry['title']
        entry['body_en'] = ''
        entry['title_cn'] = ''
        entry['body_cn'] = ''

    # 确保每个分类至少有一条
    rss_entries = ensure_min_per_category(rss_entries)

    # 按分类统计
    cat_counts = {}
    for e in rss_entries:
        cat_counts[e['category']] = cat_counts.get(e['category'], 0) + 1
    for cat, cnt in cat_counts.items():
        print(f"  {cat}: {cnt} 条")

    # 选择要抓取全文的新闻：每个分类最多取N条，共最多M条
    selected = []
    for category in ['政策/宏观经济', '行业动态', '公司新闻', '其他']:
        cat_items = [e for e in rss_entries if e['category'] == category]
        cat_items.sort(key=lambda x: x.get('published_dt', datetime.min), reverse=True)
        selected.extend(cat_items[:args.max_per_category])

    if len(selected) > args.max_total:
        selected = selected[:args.max_total]
    print(f"  已选择 {len(selected)} 条新闻进行全文抓取和翻译")

    # ── 第3步：抓取全文 ──
    print(f"\n▶ 第3步: 抓取新闻全文...")
    for i, item in enumerate(selected, 1):
        link = item.get('link', '')
        if not link:
            continue
        title_short = item['title_en'][:50]
        print(f"  [{i}/{len(selected)}] 正在抓取: {title_short}...")
        body = scrape_article_body(link, rss_summary=item.get('rss_summary', ''))
        item['body_en'] = body
        if body:
            print(f"           ✓ 获取到 {len(body)} 字符")
        else:
            print(f"           ✗ 正文获取失败")
        time.sleep(0.5)

    # ── 第4步：翻译为中文 ──
    print(f"\n▶ 第4步: 翻译标题和全文为中文...")
    for i, item in enumerate(selected, 1):
        title_short = item['title_en'][:50]
        print(f"  [{i}/{len(selected)}] 翻译中: {title_short}...")

        # 翻译标题
        item['title_cn'] = translate_text(item['title_en'])
        print(f"           标题 → {item['title_cn'][:50]}")

        # 翻译正文
        if item['body_en']:
            item['body_cn'] = translate_text(item['body_en'])
            print(f"           正文 → {len(item['body_cn'])} 字符")
        else:
            item['body_cn'] = ''

        time.sleep(0.3)

    # ── 第5步：生成报告并保存 ──
    print(f"\n▶ 第5步: 生成报告并保存...")
    save_report(selected, output_dir=args.output_dir, save_text_json=(args.report_format == 'all'))

    # ── 打印摘要 ──
    print('\n' + '=' * 70)
    print('📋 新闻摘要:')
    print('=' * 70)
    print(format_summary(selected))

    # ── 打印详细报告 ──
    print('\n' + '=' * 70)
    print('📄 详细报告:')
    print('=' * 70)
    print(format_report(selected))

    # ── 推送 ──
    if args.no_push:
        print("\n▶ 已跳过推送（--no-push）")
    else:
        print("\n▶ 推送通知...")
        pusher = NewsPusher(CONFIG_PATH)
        results = pusher.push_all(format_summary(selected))
        if results:
            print(f"  推送结果: {results}")
        else:
            print("  未配置推送渠道，仅保存到文件")

    print("\n✅ 新闻推送任务完成!")


if __name__ == '__main__':
    main()
